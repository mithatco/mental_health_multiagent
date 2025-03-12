import os
from pathlib import Path
import PyPDF2
import docx
import re
import json
from typing import List, Dict, Any, Optional, Union

class Document:
    """Represents a document with content and metadata."""
    
    def __init__(self, content: str, metadata: Dict[str, Any]):
        self.content = content
        self.metadata = metadata
    
    def __repr__(self):
        return f"Document(metadata={self.metadata}, content={self.content[:50]}...)"

class DocumentProcessor:
    """Process and extract text from various document types."""
    
    @staticmethod
    def load_document(file_path: str) -> Optional[Document]:
        """
        Load a document from a file path.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Document object or None if file type is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        # Get basic metadata
        metadata = {
            "source": file_path,
            "filename": os.path.basename(file_path),
            "extension": file_ext,
            "created": os.path.getctime(file_path),
            "modified": os.path.getmtime(file_path),
        }
        
        # Process different file types
        if file_ext == '.pdf':
            content = DocumentProcessor._extract_from_pdf(file_path)
        elif file_ext == '.txt':
            content = DocumentProcessor._extract_from_txt(file_path)
        elif file_ext in ('.docx', '.doc'):
            content = DocumentProcessor._extract_from_docx(file_path)
        elif file_ext == '.json':
            content, meta = DocumentProcessor._extract_from_json(file_path)
            metadata.update(meta)
        else:
            print(f"Unsupported file type: {file_ext}")
            return None
        
        if not content:
            print(f"Warning: No content extracted from {file_path}")
            
        return Document(content, metadata)
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {str(e)}")
        
        return text
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            content = [paragraph.text for paragraph in doc.paragraphs]
            return '\n'.join(content)
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_from_json(file_path: str) -> tuple:
        """Extract text and metadata from a JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Assuming JSON has 'content' field, or convert the whole JSON to string
            content = data.get('content', json.dumps(data))
            
            # Extract any additional metadata
            metadata = {k: v for k, v in data.items() if k != 'content'}
            
            return content, metadata
        except Exception as e:
            print(f"Error extracting text from JSON {file_path}: {str(e)}")
            return "", {}

class Chunker:
    """Split documents into smaller chunks for processing."""
    
    @staticmethod
    def chunk_document(document: Document, 
                      chunk_size: int = 1000, 
                      overlap: int = 200) -> List[Document]:
        """
        Split a document into overlapping chunks.
        
        Args:
            document: Document object to chunk
            chunk_size: Maximum size of each chunk in characters
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of Document objects (chunks)
        """
        content = document.content
        chunks = []
        
        # For very short documents, don't chunk
        if len(content) <= chunk_size:
            return [document]
        
        # Split into chunks
        start = 0
        chunk_id = 0
        
        while start < len(content):
            # Calculate end position
            end = min(start + chunk_size, len(content))
            
            # If not the first chunk, include overlap
            if start > 0:
                start = max(0, start - overlap)
            
            # Try to end chunks at sentence boundaries when possible
            if end < len(content):
                # Look for sentence end within the last 100 chars of the chunk
                search_zone = content[max(end-100, 0):end]
                sentence_ends = [m.end() + max(end-100, 0) for m in re.finditer(r'[.!?]\s+', search_zone)]
                
                if sentence_ends:
                    # Use the last sentence end as the chunk end
                    end = sentence_ends[-1]
            
            # Create chunk with same metadata + chunk-specific metadata
            chunk_metadata = document.metadata.copy()
            chunk_metadata.update({
                "chunk_id": chunk_id,
                "chunk_start": start,
                "chunk_end": end,
                "is_chunk": True
            })
            
            chunk_document = Document(content[start:end], chunk_metadata)
            chunks.append(chunk_document)
            
            # Move to next chunk
            start = end
            chunk_id += 1
        
        return chunks

def extract_questions_from_text(text):
    """Extract questions from text content."""
    if not text:
        return []
    
    # Preprocess text to handle common issues
    processed_text = text
    
    # Handle common questionnaire formatting patterns
    # Replace line breaks within numbered items to help with pattern matching
    processed_text = re.sub(r'(\d+\.?\s*[A-Za-z][^?]*?)\n\s*([^A-Z0-9\n][^?]*?\?)', r'\1 \2', processed_text)
    
    # Handle multi-line questions with parenthetical content
    processed_text = re.sub(r'\(e\.g\.,\s*([^\)]*?)\n\s*([^\)]*?)\)', r'(e.g., \1 \2)', processed_text)
    processed_text = re.sub(r'\(([^\)]*?)\n\s*([^\)]*?)\)', r'(\1 \2)', processed_text)
    
    # Replace common line break patterns within questions with spaces
    processed_text = re.sub(r'(\S)\n\s*(\S)', r'\1 \2', processed_text)
    
    # Replace multiple spaces with a single space
    processed_text = re.sub(r'\s+', ' ', processed_text)
    
    # Special handling for DSM-5 questionnaire format
    dsm5_questions = []
    
    # First try to extract questions from DSM-5 format (numbered list with question marks)
    dsm5_pattern = r'(\d+\.?\s*[A-Za-z][^.!?]*?\?)'
    dsm5_matches = re.findall(dsm5_pattern, processed_text)
    
    if dsm5_matches and len(dsm5_matches) >= 15:  # DSM-5 usually has many questions
        dsm5_questions = dsm5_matches
    
    # Look for questions containing "e.g." or parenthetical content
    complex_pattern = r'([A-Za-z][^.!?]*?e\.g\.,.*?\)?[^.!?]*?\?)'
    complex_matches = re.findall(complex_pattern, processed_text)
    
    # Merge the results, prioritizing more specific matches
    if complex_matches:
        # Map to identify duplicates
        question_map = {}
        for q in dsm5_questions:
            # Extract just the start of the question to use as a key
            key = re.sub(r'^(\d+\.?\s*)', '', q.strip()).lower()[:30]
            question_map[key] = q
            
        # Add complex matches if they're not duplicates
        for q in complex_matches:
            key = q.lower()[:30]
            if key not in question_map:
                question_map[key] = q
                
        dsm5_questions = list(question_map.values())
    
    # If DSM-5 specific extraction worked well, use those results
    if len(dsm5_questions) >= 15:
        questions = dsm5_questions
    else:
        # Otherwise fall back to generic question extraction
        # Pattern 1: Look for sentences ending with question marks
        questions = re.findall(r'([A-Z][^.!?]*?\?)', processed_text)
        
        # If the above pattern doesn't find enough questions, try alternative patterns
        if len(questions) < 10:  # Threshold for a typical questionnaire
            # Pattern 2: Look for numbered questions
            numbered_questions = re.findall(r'(?:\d+\.?\s*)([A-Za-z][^.!?]*?\?)', processed_text)
            if len(numbered_questions) > len(questions):
                questions = numbered_questions
        
        # Pattern 3: Look for questions in a questionnaire format with possible multi-line content
        if len(questions) < 10:
            complex_questions = []
            lines = text.split('\n')
            question_buffer = ""
            in_question = False
            
            for line in lines:
                clean_line = line.strip()
                if not clean_line:
                    if question_buffer and ('?' in question_buffer or in_question):
                        complex_questions.append(question_buffer.strip())
                        question_buffer = ""
                        in_question = False
                    continue
                
                # Check if this line is likely part of a question
                has_question_marker = '?' in clean_line
                is_numbered_item = re.match(r'^\d+\.?\s*[A-Z]', clean_line)
                continues_previous = not clean_line[0].isupper() and not is_numbered_item
                
                if is_numbered_item:
                    # If we already have a question in buffer, save it
                    if question_buffer and (in_question or '?' in question_buffer):
                        complex_questions.append(question_buffer.strip())
                    question_buffer = clean_line
                    in_question = True
                elif continues_previous and (in_question or question_buffer):
                    # This line continues a previous question
                    question_buffer += " " + clean_line
                    if has_question_marker:
                        in_question = False  # Question is complete
                elif has_question_marker or clean_line.endswith(':'):
                    # This might be a standalone question
                    if question_buffer and (in_question or '?' in question_buffer):
                        complex_questions.append(question_buffer.strip())
                    question_buffer = clean_line
                    in_question = not has_question_marker
                else:
                    # Not sure what this is - if it starts with a capital, treat as new item
                    if clean_line[0].isupper():
                        if question_buffer and (in_question or '?' in question_buffer):
                            complex_questions.append(question_buffer.strip())
                            question_buffer = ""
                            in_question = False
                        question_buffer = clean_line
                    else:
                        question_buffer += " " + clean_line
            
            # Don't forget the last question
            if question_buffer and (in_question or '?' in question_buffer):
                complex_questions.append(question_buffer.strip())
                
            if len(complex_questions) > len(questions):
                questions = complex_questions
    
    # Final processing to clean up questions
    cleaned_questions = []
    for q in questions:
        # Remove leading numbers and whitespace
        q = re.sub(r'^(\d+\.?\s*)', '', q.strip())
        
        # Make sure it starts with a capital letter
        if q and not q[0].isupper():
            q = q[0].upper() + q[1:]
            
        # Make sure it ends with a question mark
        if q and not q.endswith('?'):
            q += '?'
            
        # Validate that it's actually a question and not just a fragment
        if q and len(q) > 10 and '?' in q:
            cleaned_questions.append(q)
    
    return cleaned_questions

def process_documents_directory(directory_path: str) -> Dict[str, List[Document]]:
    """
    Process all documents in a directory.
    
    Args:
        directory_path: Path to the directory containing documents
        
    Returns:
        Dictionary mapping file types to lists of Document objects
    """
    document_map = {}
    print(f"Scanning directory: {directory_path}")
    
    try:
        files = [os.path.join(directory_path, f) for f in os.listdir(directory_path) 
                if os.path.isfile(os.path.join(directory_path, f))]
        
        print(f"Found {len(files)} files in directory")
        
        for file_path in files:
            print(f"Processing file: {file_path}")
            document = DocumentProcessor.load_document(file_path)
            if document:
                ext = Path(file_path).suffix.lower()
                if ext not in document_map:
                    document_map[ext] = []
                document_map[ext].append(document)
                print(f"Added document with ext {ext}, content length: {len(document.content)}")
            else:
                print(f"Failed to load document: {file_path}")
    
    except Exception as e:
        print(f"Error processing directory {directory_path}: {str(e)}")
    
    return document_map

def get_questionnaires_from_documents(documents: List[Document]) -> Dict[str, List[str]]:
    """
    Extract questionnaires from documents.
    
    Args:
        documents: List of Document objects
        
    Returns:
        Dictionary mapping document names to lists of questions
    """
    questionnaire_map = {}
    
    for document in documents:
        questions = extract_questions_from_text(document.content)
        if questions:
            name = document.metadata.get("filename", "Unknown")
            questionnaire_map[name] = questions
    
    return questionnaire_map
